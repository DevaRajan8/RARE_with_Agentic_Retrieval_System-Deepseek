import streamlit as st
import requests
import json
import time
import logging
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
import numpy as np
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
import hashlib
import os
from datetime import datetime
import traceback
import PyPDF2
import docx
from docx import Document
import pandas as pd
from io import BytesIO, StringIO
import csv
from pathlib import Path
import zipfile

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class RetrievedDocument:
    """Represents a retrieved document with metadata."""
    content: str
    similarity_score: float
    source: str
    timestamp: str

class DocumentProcessor:
    """Handles processing of multiple document formats."""
    
    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from TXT files."""
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_content.decode('latin-1')
            except:
                return file_content.decode('utf-8', errors='ignore')
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF files."""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return f"Error processing PDF: {str(e)}"
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX files."""
        try:
            doc_file = BytesIO(file_content)
            doc = Document(doc_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return f"Error processing DOCX: {str(e)}"
    
    @staticmethod
    def extract_text_from_csv(file_content: bytes) -> str:
        """Extract text from CSV files."""
        try:
            csv_content = file_content.decode('utf-8')
            csv_file = StringIO(csv_content)
            reader = csv.reader(csv_file)
            text = ""
            for row_num, row in enumerate(reader):
                if row_num == 0:  # Header row
                    text += "Headers: " + ", ".join(row) + "\n\n"
                else:
                    text += "Row " + str(row_num) + ": " + ", ".join(row) + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from CSV: {e}")
            return f"Error processing CSV: {str(e)}"
    
    @staticmethod
    def extract_text_from_json(file_content: bytes) -> str:
        """Extract text from JSON files."""
        try:
            json_content = file_content.decode('utf-8')
            data = json.loads(json_content)
            def json_to_text(obj, prefix=""):
                text = ""
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            text += f"{prefix}{key}:\n"
                            text += json_to_text(value, prefix + "  ")
                        else:
                            text += f"{prefix}{key}: {value}\n"
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        text += f"{prefix}Item {i+1}:\n"
                        text += json_to_text(item, prefix + "  ")
                else:
                    text += f"{prefix}{obj}\n"
                return text
            return json_to_text(data).strip()
        except Exception as e:
            logger.error(f"Error extracting text from JSON: {e}")
            return f"Error processing JSON: {str(e)}"
    
    @staticmethod
    def extract_text_from_excel(file_content: bytes) -> str:
        """Extract text from Excel files (XLS, XLSX)."""
        try:
            excel_file = BytesIO(file_content)
            try:
                df_dict = pd.read_excel(excel_file, sheet_name=None)
            except:
                df_dict = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
            text = ""
            for sheet_name, df in df_dict.items():
                text += f"Sheet: {sheet_name}\n"
                text += f"Headers: {', '.join(df.columns.astype(str))}\n\n"
                for index, row in df.iterrows():
                    row_text = [str(row[col]) if pd.notna(row[col]) else "" for col in df.columns]
                    text += f"Row {index + 1}: {', '.join(row_text)}\n"
                text += "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from Excel: {e}")
            return f"Error processing Excel file: {str(e)}"
    
    @classmethod
    def process_uploaded_file(cls, uploaded_file) -> Tuple[str, str]:
        """Process an uploaded file and extract text content."""
        file_extension = Path(uploaded_file.name).suffix.lower()
        file_content = uploaded_file.read()
        if file_extension == '.txt':
            text = cls.extract_text_from_txt(file_content)
        elif file_extension == '.pdf':
            text = cls.extract_text_from_pdf(file_content)
        elif file_extension == '.docx':
            text = cls.extract_text_from_docx(file_content)
        elif file_extension == '.csv':
            text = cls.extract_text_from_csv(file_content)
        elif file_extension == '.json':
            text = cls.extract_text_from_json(file_content)
        elif file_extension in ['.xlsx', '.xls']:
            text = cls.extract_text_from_excel(file_content)
        else:
            text = cls.extract_text_from_txt(file_content)
        return text, uploaded_file.name

class EnhancedRAGSystem:
    """Enhanced RAG system with chain-of-thought reasoning for DeepSeek model."""
    
    def __init__(self, groq_api_key: str):
        self.groq_api_key = groq_api_key
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.vector_db = None
        self.collection_name = "enhanced_rag_collection"
        self.db_path = "./chroma_db"
        self._initialize_vector_db()
        
    def _initialize_vector_db(self):
        """Initialize ChromaDB with persistent storage."""
        try:
            client = chromadb.PersistentClient(
                path=self.db_path,
                settings=Settings(anonymized_telemetry=False)
            )
            try:
                self.vector_db = client.get_collection(name=self.collection_name)
                logger.info(f"Loaded existing collection with {self.vector_db.count()} documents")
            except:
                self.vector_db = client.create_collection(
                    name=self.collection_name,
                    metadata={"hnsw:space": "cosine"}
                )
                logger.info("Created new vector database collection")
        except Exception as e:
            logger.error(f"Failed to initialize vector database: {e}")
            raise
        
    @staticmethod
    def initialize_rag_system():
        """Initialize RAG system with proper session state management."""
        groq_api_key = os.getenv("GROQ_API_KEY")
        if not groq_api_key:
            st.warning("Please set your GROQ_API_KEY environment variable.")
            st.stop()
        if 'rag_system' not in st.session_state:
            try:
                with st.spinner("Initializing RAG system..."):
                    st.session_state.rag_system = EnhancedRAGSystem(groq_api_key)
                doc_count = st.session_state.rag_system.get_document_count()
                if doc_count > 0:
                    st.success(f"RAG system initialized with {doc_count} existing documents!")
                else:
                    st.success("RAG system initialized!")
            except Exception as e:
                st.error(f"Failed to initialize RAG system: {e}")
                st.stop()
        
    def clear_knowledge_base(self):
        """Clear all documents from the knowledge base."""
        try:
            client = chromadb.PersistentClient(
                path=self.db_path,
                settings=Settings(anonymized_telemetry=False)
            )
            try:
                client.delete_collection(name=self.collection_name)
            except:
                pass
            self.vector_db = client.create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            logger.info("Knowledge base cleared successfully")
            return True
        except Exception as e:
            logger.error(f"Failed to clear knowledge base: {e}")
            return False

    def add_documents(self, documents: List[str], sources: List[str] = None):
        """Add documents to the vector database with embeddings."""
        if not documents:
            return
        if sources is None:
            sources = [f"doc_{i}" for i in range(len(documents))]
        try:
            embeddings = self.embedding_model.encode(documents).tolist()
            ids = [hashlib.md5(doc.encode()).hexdigest() for doc in documents]
            self.vector_db.add(
                embeddings=embeddings,
                documents=documents,
                metadatas=[{"source": source, "timestamp": datetime.now().isoformat()} 
                          for source in sources],
                ids=ids
            )
            logger.info(f"Added {len(documents)} documents to vector database")
        except Exception as e:
            logger.error(f"Failed to add documents: {e}")
            raise

    def retrieve_relevant_documents(self, query: str, top_k: int = 5) -> List[RetrievedDocument]:
        """Retrieve most relevant documents using semantic similarity."""
        try:
            query_embedding = self.embedding_model.encode([query]).tolist()[0]
            results = self.vector_db.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                include=['documents', 'metadatas', 'distances']
            )
            retrieved_docs = []
            for i in range(len(results['documents'][0])):
                doc = RetrievedDocument(
                    content=results['documents'][0][i],
                    similarity_score=1 - results['distances'][0][i],
                    source=results['metadatas'][0][i].get('source', 'unknown'),
                    timestamp=results['metadatas'][0][i].get('timestamp', 'unknown')
                )
                retrieved_docs.append(doc)
            logger.info(f"Retrieved {len(retrieved_docs)} relevant documents")
            return retrieved_docs
        except Exception as e:
            logger.error(f"Failed to retrieve documents: {e}")
            return []
        
    def get_document_count(self):
        """Get the number of documents in the knowledge base."""
        try:
            if self.vector_db:
                return self.vector_db.count()
            return 0
        except:
            return 0

    @staticmethod
    def display_system_stats():
        """Display system statistics in sidebar."""
        if 'rag_system' in st.session_state:
            try:
                doc_count = st.session_state.rag_system.get_document_count()
                st.metric("Documents in Knowledge Base", doc_count)
                if doc_count > 0:
                    st.caption("📅 Knowledge base persisted across sessions")
                else:
                    st.caption("📝 Add documents to build your knowledge base")
            except Exception as e:
                st.error(f"Error getting system stats: {e}")

    @staticmethod
    def handle_clear_knowledge_base():
        """Handle clearing the knowledge base with confirmation."""
        if st.button("🗑️ Clear Knowledge Base", type="secondary", key="clear_kb_button"):
            if 'confirm_clear' not in st.session_state:
                st.session_state.confirm_clear = False
            if not st.session_state.confirm_clear:
                st.session_state.confirm_clear = True
                st.warning("⚠️ This will delete all documents. Click again to confirm.")
            else:
                try:
                    success = st.session_state.rag_system.clear_knowledge_base()
                    if success:
                        st.success("Knowledge base cleared!")
                        st.session_state.confirm_clear = False
                        st.rerun()
                    else:
                        st.error("Failed to clear knowledge base")
                except Exception as e:
                    st.error(f"Error clearing knowledge base: {e}")
        if st.session_state.get('confirm_clear', False):
            if st.button("Cancel", key="cancel_clear_kb"):
                st.session_state.confirm_clear = False
                st.rerun()

    def call_groq_api(self, messages: list, max_retries: int = 3) -> str:
        """Call Groq API with robust error handling and retry mechanism."""
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {self.groq_api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "deepseek-r1-distill-llama-70b",
            "messages": messages,
            "max_tokens": 1500,
            "temperature": 0.3
        }
        for attempt in range(max_retries):
            try:
                response = requests.post(url, headers=headers, json=payload, timeout=30)
                if response.status_code == 200:
                    data = response.json()
                    return data['choices'][0]['message']['content']
                elif response.status_code == 429:
                    wait_time = 2 ** attempt
                    logger.warning(f"Rate limited. Waiting {wait_time} seconds...")
                    time.sleep(wait_time)
                    continue
                else:
                    logger.error(f"API error {response.status_code}: {response.text}")
                    if attempt == max_retries - 1:
                        raise Exception(f"API call failed: {response.status_code}")
            except requests.exceptions.Timeout:
                logger.warning(f"Request timeout on attempt {attempt + 1}")
                if attempt == max_retries - 1:
                    raise Exception("Request timed out after all retries")
            except Exception as e:
                logger.error(f"API call error: {e}")
                if attempt == max_retries - 1:
                    raise
            time.sleep(1)
        raise Exception("Maximum retries exceeded")

    def generate_chain_of_thought_prompt(self, query: str, retrieved_docs: List[RetrievedDocument]) -> str:
        """Generate an enhanced chain-of-thought prompt for reasoning."""
        context_sections = [f"""
Document {i} (Source: {doc.source}, Similarity: {doc.similarity_score:.3f}):
{doc.content}
""" for i, doc in enumerate(retrieved_docs, 1)]
        context = "\n".join(context_sections)
        chain_of_thought_prompt = f"""You are an advanced AI assistant with enhanced reasoning capabilities. You will answer the user's question using a systematic chain-of-thought approach.

RETRIEVED CONTEXT:
{context}

USER QUESTION: {query}

INSTRUCTIONS:
Please follow this systematic reasoning process and reply with a well-reasoned answer , you can use the retrieved documents to support your reasoning and you should use these steps for replying for the answer also:

1. UNDERSTANDING: First, carefully analyze and restate the user's question to ensure clear understanding.

2. CONTEXT ANALYSIS: Examine the retrieved documents and identify:
   - Key relevant information
   - Important facts and details
   - Any contradictions or gaps
   - The reliability and relevance of each source

3. REASONING PROCESS: Think through the problem step-by-step:
   - Break down complex questions into smaller components
   - Apply logical reasoning to connect information
   - Consider multiple perspectives or approaches
   - Identify any assumptions being made

4. EVIDENCE SYNTHESIS: Combine information from multiple sources:
   - Prioritize the most relevant and reliable information
   - Resolve any contradictions found
   - Build a coherent understanding

5. CONCLUSION FORMATION: Based on your analysis:
   - Provide a clear, well-reasoned answer
   - Explain your reasoning process
   - Acknowledge any limitations or uncertainties
   - Suggest follow-up questions if relevant

6. VERIFICATION: Double-check your answer:
   - Ensure it directly addresses the user's question
   - Verify consistency with the evidence
   - Consider alternative interpretations

Please structure your response clearly, showing each step of your reasoning process. Be thorough but concise, and always ground your answer in the provided context while being transparent about your reasoning.

Begin your systematic analysis now:"""
        return chain_of_thought_prompt

    def generate_enhanced_response(self, query: str, top_k: int = 5) -> Dict[str, Any]:
        """Generate an enhanced response using chain-of-thought reasoning."""
        start_time = time.time()
        try:
            retrieved_docs = self.retrieve_relevant_documents(query, top_k)
            if not retrieved_docs:
                return {
                    "answer": "I apologize, but I couldn't find relevant documents to answer your question. Please try rephrasing your query or add more documents to the knowledge base.",
                    "retrieved_documents": [],
                    "reasoning_steps": [],
                    "processing_time": time.time() - start_time,
                    "success": False
                }
            cot_prompt = self.generate_chain_of_thought_prompt(query, retrieved_docs)
            messages = [
                {"role": "system", "content": "You are an expert AI assistant specializing in systematic reasoning and analysis."},
                {"role": "user", "content": cot_prompt}
            ]
            response = self.call_groq_api(messages)
            reasoning_steps = self._extract_reasoning_steps(response)
            processing_time = time.time() - start_time
            return {
                "answer": response,
                "retrieved_documents": [
                    {
                        "content": doc.content[:200] + "..." if len(doc.content) > 200 else doc.content,
                        "source": doc.source,
                        "similarity_score": doc.similarity_score,
                        "timestamp": doc.timestamp
                    } for doc in retrieved_docs
                ],
                "reasoning_steps": reasoning_steps,
                "processing_time": processing_time,
                "success": True
            }
        except Exception as e:
            logger.error(f"Error generating response: {e}")
            return {
                "answer": f"An error occurred while processing your request: {str(e)}",
                "retrieved_documents": [],
                "reasoning_steps": [],
                "processing_time": time.time() - start_time,
                "success": False
            }

    def _extract_reasoning_steps(self, response: str) -> List[str]:
        """Extract reasoning steps from the response."""
        steps = []
        lines = response.split('\n')
        current_step = ""
        step_indicators = ['1.', '2.', '3.', '4.', '5.', '6.', 'UNDERSTANDING:', 'CONTEXT ANALYSIS:', 
                          'REASONING PROCESS:', 'EVIDENCE SYNTHESIS:', 'CONCLUSION FORMATION:', 'VERIFICATION:']
        for line in lines:
            line = line.strip()
            if any(indicator in line for indicator in step_indicators):
                if current_step:
                    steps.append(current_step.strip())
                current_step = line
            elif line and current_step:
                current_step += " " + line
        if current_step:
            steps.append(current_step.strip())
        return steps[:6]

def main():
    st.set_page_config(
        page_title="Enhanced RAG with Chain-of-Thought Reasoning",
        page_icon="🧠",
        layout="wide"
    )
    EnhancedRAGSystem.initialize_rag_system()
    
    st.title("🧠 Enhanced RAG with Chain-of-Thought Reasoning")
    st.markdown("*Powered by DeepSeek with Advanced Retrieval and Reasoning*")
    
    with st.sidebar:
        st.header("⚙️ Configuration")
        groq_api_key = os.getenv("GROQ_API_KEY")
        if not groq_api_key:
            st.warning("Please enter your Groq API key to continue.")
            st.stop()
        st.divider()
        st.subheader("📊 System Stats")
        EnhancedRAGSystem.display_system_stats()
        st.subheader("📁 Supported Formats")
        supported_formats = {
            "📄 Text": "TXT files",
            "📕 PDF": "PDF documents",
            "📘 Word": "DOCX files",
            "📊 Excel": "XLSX, XLS files",
            "📋 CSV": "Comma-separated values",
            "🔧 JSON": "JSON data files"
        }
        for icon_format, description in supported_formats.items():
            st.text(f"{icon_format}: {description}")
        st.caption("Upload any combination of these file types for comprehensive document analysis.")
        st.subheader("🔍 Retrieval Settings")
        top_k = st.slider("Number of documents to retrieve", 1, 10, 5)
        st.divider()
        st.subheader("📊 System Stats")
        if hasattr(st.session_state.rag_system, 'vector_db') and st.session_state.rag_system.vector_db:
            try:
                doc_count = st.session_state.rag_system.vector_db.count()
                st.metric("Documents in Knowledge Base", doc_count)
            except:
                st.metric("Documents in Knowledge Base", "Unknown")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📚 Knowledge Base Management")
        uploaded_files = st.file_uploader(
            "Upload documents (Multiple formats supported)",
            type=['txt', 'pdf', 'docx', 'csv', 'json', 'xlsx', 'xls'],
            accept_multiple_files=True,
            help="Supported formats: TXT, PDF, DOCX, CSV, JSON, XLSX, XLS"
        )
        with st.expander("📝 Add Documents Manually"):
            manual_docs = st.text_area(
                "Enter documents (separate multiple documents with '---')",
                height=150,
                placeholder="Enter your documents here. Use '---' to separate multiple documents."
            )
            manual_sources = st.text_input(
                "Document sources (comma-separated)",
                placeholder="source1, source2, source3..."
            )
            if st.button("Add Manual Documents", key="add_manual_docs"):
                if manual_docs.strip():
                    docs = [doc.strip() for doc in manual_docs.split('---') if doc.strip()]
                    sources = [s.strip() for s in manual_sources.split(',')] if manual_sources else None
                    if sources and len(sources) != len(docs):
                        st.warning("Number of sources must match number of documents")
                    else:
                        try:
                            with st.spinner("Adding documents..."):
                                st.session_state.rag_system.add_documents(docs, sources)
                            st.success(f"Added {len(docs)} documents to knowledge base!")
                        except Exception as e:
                            st.error(f"Failed to add documents: {e}")
        if uploaded_files:
            docs_to_add = []
            sources_to_add = []
            processing_status = []
            progress_bar = st.progress(0)
            status_text = st.empty()
            for i, file in enumerate(uploaded_files):
                try:
                    status_text.text(f"Processing {file.name}...")
                    progress_bar.progress((i + 1) / len(uploaded_files))
                    text_content, source_name = DocumentProcessor.process_uploaded_file(file)
                    if text_content and not text_content.startswith("Error"):
                        docs_to_add.append(text_content)
                        sources_to_add.append(source_name)
                        processing_status.append(f"✅ {file.name}: Successfully processed")
                    else:
                        processing_status.append(f"❌ {file.name}: {text_content}")
                except Exception as e:
                    processing_status.append(f"❌ {file.name}: Error - {str(e)}")
            progress_bar.empty()
            status_text.empty()
            if processing_status:
                with st.expander("📋 File Processing Results", expanded=True):
                    for status in processing_status:
                        if status.startswith("✅"):
                            st.success(status)
                        else:
                            st.error(status)
            if docs_to_add:
                try:
                    with st.spinner("Adding processed documents to knowledge base..."):
                        st.session_state.rag_system.add_documents(docs_to_add, sources_to_add)
                    st.success(f"Successfully added {len(docs_to_add)} documents to knowledge base!")
                except Exception as e:
                    st.error(f"Failed to add documents to knowledge base: {e}")
            else:
                if uploaded_files:
                    st.warning("No documents could be processed successfully.")
    
    with col2:
        st.header("🔧 Quick Actions")
        if st.button("📚 Load Sample Documents", key="load_samples"):
            sample_docs = [
                "Artificial Intelligence (AI) is a branch of computer science that aims to create intelligent machines that work and react like humans. AI research has been highly successful in developing effective techniques for solving a wide range of problems, from game playing to medical diagnosis. Modern AI systems use machine learning, deep learning, and neural networks to process vast amounts of data and make intelligent decisions.",
                "Machine Learning Pipeline: 1. Data Collection - Gather relevant datasets from various sources. 2. Data Preprocessing - Clean, normalize, and transform data. 3. Feature Engineering - Select and create meaningful features. 4. Model Selection - Choose appropriate algorithms. 5. Training - Fit the model to training data. 6. Evaluation - Assess model performance using metrics. 7. Deployment - Implement model in production environment.",
                "Deep Learning for Computer Vision: A Comprehensive Survey. Abstract: This paper presents a comprehensive review of deep learning techniques applied to computer vision tasks. We analyze convolutional neural networks (CNNs), their architectures including ResNet, VGG, and Inception, and their applications in image classification, object detection, and semantic segmentation. Our study covers recent advances in attention mechanisms, transformer architectures, and their impact on visual recognition tasks.",
                "Company Policy: Remote Work Guidelines. Effective Date: January 2024. Purpose: To establish clear guidelines for remote work arrangements. Eligibility: All full-time employees who have completed probationary period. Requirements: Reliable internet connection (minimum 25 Mbps), dedicated workspace, availability during core hours (9 AM - 3 PM local time). Communication: Daily check-ins via Slack, weekly video conferences, monthly in-person meetings.",
                "Q4 2023 Sales Analysis: Total Revenue: $2.5M (15% increase from Q3). Top Performing Products: Product A (35% of sales), Product B (28% of sales), Product C (22% of sales). Geographic Distribution: North America (45%), Europe (30%), Asia-Pacific (25%). Customer Acquisition Cost: $125 (down 8% from previous quarter). Customer Lifetime Value: $1,850 (up 12%). Recommendations: Increase marketing budget for Product A, expand European operations, improve customer retention programs.",
                "API Documentation: Authentication Endpoint. URL: /api/v1/auth/login. Method: POST. Headers: Content-Type: application/json. Payload: {username: string, password: string}. Response: {token: string, expires_in: number, user_id: string}. Error Codes: 400 (Invalid request), 401 (Invalid credentials), 429 (Rate limit exceeded), 500 (Server error). Rate Limit: 5 requests per minute per IP address."
            ]
            sample_sources = [
                "AI_Overview_2024",
                "ML_Pipeline_Guide",
                "Computer_Vision_Survey",
                "Remote_Work_Policy",
                "Q4_Sales_Report",
                "API_Documentation"
            ]
            try:
                with st.spinner("Loading sample documents..."):
                    st.session_state.rag_system.add_documents(sample_docs, sample_sources)
                st.success("Enhanced sample documents loaded!")
                st.info("Sample documents include: AI/ML content, technical documentation, research abstracts, business policies, data reports, and API specifications.")
            except Exception as e:
                st.error(f"Failed to load sample documents: {e}")
        EnhancedRAGSystem.handle_clear_knowledge_base()
    
    st.divider()
    
    st.header("🤔 Ask Questions")
    query = st.text_input(
        "Enter your question:",
        placeholder="What would you like to know about the documents?",
        help="Ask any question about the documents in your knowledge base",
        key="query_input"
    )
    with st.expander("💡 Example Queries"):
        st.write("Try these example questions:")
        example_queries = [
            "What is the difference between AI and Machine Learning?",
            "How does deep learning relate to neural networks?",
            "What are the steps in a typical machine learning pipeline?",
            "What are the remote work policy requirements?",
            "How did sales perform in Q4 2023?",
            "What are the API authentication requirements?",
            "Compare computer vision techniques mentioned in the research",
            "What factors contributed to the revenue increase?",
            "Explain the error codes for the authentication endpoint"
        ]
        for i, example in enumerate(example_queries):
            if st.button(f"📝 {example}", key=f"example_{i}"):
                query = example
                st.rerun()
    if st.button("🚀 Generate Enhanced Response", type="primary", disabled=not query, key="generate_response"):
        if not query.strip():
            st.warning("Please enter a question.")
        else:
            with st.spinner("🧠 Thinking with chain-of-thought reasoning..."):
                result = st.session_state.rag_system.generate_enhanced_response(query, top_k)
            if result['success']:
                st.subheader("✨ Enhanced Answer")
                st.write(result['answer'])
                tab1, tab2, tab3 = st.tabs(["🧠 Reasoning Steps", "📄 Retrieved Documents", "📊 Metrics"])
                with tab1:
                    st.subheader("Chain-of-Thought Reasoning Process")
                    if result['reasoning_steps']:
                        for i, step in enumerate(result['reasoning_steps'], 1):
                            with st.expander(f"Step {i}", expanded=True):
                                st.write(step)
                    else:
                        st.info("Reasoning steps could not be extracted from the response.")
                with tab2:
                    st.subheader("Retrieved Documents")
                    for i, doc in enumerate(result['retrieved_documents'], 1):
                        with st.expander(f"Document {i} - {doc['source']} (Similarity: {doc['similarity_score']:.3f})"):
                            st.write(doc['content'])
                            st.caption(f"Source: {doc['source']} | Timestamp: {doc['timestamp']}")
                with tab3:
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Processing Time", f"{result['processing_time']:.2f}s")
                    with col2:
                        st.metric("Documents Retrieved", len(result['retrieved_documents']))
                    with col3:
                        st.metric("Reasoning Steps", len(result['reasoning_steps']))
            else:
                st.error("Failed to generate response")
                st.write(result['answer'])

if __name__ == "__main__":
    main()