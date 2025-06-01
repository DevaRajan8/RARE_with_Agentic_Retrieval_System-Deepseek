import streamlit as st
import os
import json
import time
import requests
from typing import List, Dict, Optional
from dataclasses import dataclass
from enum import Enum
import hashlib
import pickle
from pathlib import Path
import PyPDF2
import docx
import io
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import uuid

# Configuration
class RetrievalMode(Enum):
    CHUNKS = "chunks"
    FILES_VIA_METADATA = "files_via_metadata"
    FILES_VIA_CONTENT = "files_via_content"
    WEB_SEARCH = "web_search"  # New mode for external tool use

@dataclass
class Document:
    id: str
    content: str
    metadata: Dict
    file_path: str
    chunks: List[str] = None
    embedding: np.ndarray = None

@dataclass
class RetrievalResult:
    content: str
    metadata: Dict
    score: float
    retrieval_mode: str

class DocumentProcessor:
    def __init__(self):
        try:
            self.embedder = SentenceTransformer('all-MiniLM-L6-v2')
            st.success("✅ Embedding model loaded successfully!")
        except Exception as e:
            st.error(f"Error loading embedding model: {e}")
            self.embedder = None
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text += page_text + "\n"
                st.info(f"Extracted {len(page_text)} characters from page {page_num + 1}")
            if not text.strip():
                st.warning("PDF appears to be empty or contains only images/scanned content")
            return text
        except Exception as e:
            st.error(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_content.decode('latin-1')
            except Exception as e:
                st.error(f"Error extracting TXT text: {e}")
                return ""
    
    def process_uploaded_file(self, uploaded_file) -> Document:
        try:
            file_content = uploaded_file.read()
            filename = uploaded_file.name
            file_extension = filename.split('.')[-1].lower()
            st.info(f"Processing {filename} ({len(file_content)} bytes)")
            if file_extension == 'pdf':
                text = self.extract_text_from_pdf(file_content)
            elif file_extension == 'docx':
                text = self.extract_text_from_docx(file_content)
            elif file_extension == 'txt':
                text = self.extract_text_from_txt(file_content)
            else:
                st.error(f"Unsupported file type: {file_extension}")
                return None
            if not text.strip():
                st.error(f"No text extracted from {filename}")
                return None
            st.success(f"✅ Extracted {len(text)} characters from {filename}")
            chunks = self.create_chunks(text)
            st.info(f"Created {len(chunks)} chunks")
            embedding = None
            if self.embedder:
                try:
                    embedding_text = text[:1000] if len(text) > 1000 else text
                    embedding = self.embedder.encode(embedding_text)
                    st.success("✅ Created document embedding")
                except Exception as e:
                    st.warning(f"Could not create embedding: {e}")
            document = Document(
                id=str(uuid.uuid4()),
                content=text,
                metadata={
                    'filename': filename,
                    'file_type': file_extension,
                    'size': len(text),
                    'chunk_count': len(chunks),
                    'word_count': len(text.split())
                },
                file_path=filename,
                chunks=chunks,
                embedding=embedding
            )
            return document
        except Exception as e:
            st.error(f"Error processing file {uploaded_file.name}: {e}")
            return None
    
    def create_chunks(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        words = text.split()
        chunks = []
        if len(words) <= chunk_size:
            chunks.append(text)
        else:
            for i in range(0, len(words), chunk_size - overlap):
                chunk_words = words[i:i + chunk_size]
                chunk = ' '.join(chunk_words)
                if chunk.strip():
                    chunks.append(chunk)
        return chunks

class KnowledgeBase:
    def __init__(self, name: str, description: str, documents: List[Document] = None):
        self.name = name
        self.description = description
        self.documents = documents or []
        self.processor = DocumentProcessor()
    
    def add_document(self, document: Document):
        self.documents.append(document)
        st.success(f"Added document: {document.metadata['filename']} to KB: {self.name}")
    
    def get_relevant_chunks(self, query: str, top_k: int = 5) -> List[RetrievalResult]:
        results = []
        st.info(f"Searching {len(self.documents)} documents with {sum(len(doc.chunks or []) for doc in self.documents)} total chunks")
        query_embedding = None
        if self.processor.embedder:
            try:
                query_embedding = self.processor.embedder.encode(query)
                st.info("✅ Created query embedding")
            except Exception as e:
                st.warning(f"Could not create query embedding: {e}")
        for doc in self.documents:
            for i, chunk in enumerate(doc.chunks or []):
                if query_embedding is not None and self.processor.embedder:
                    try:
                        chunk_embedding = self.processor.embedder.encode(chunk)
                        score = cosine_similarity([query_embedding], [chunk_embedding])[0][0]
                    except:
                        score = self._calculate_keyword_relevance(query, chunk)
                else:
                    score = self._calculate_keyword_relevance(query, chunk)
                if score > 0.05:
                    results.append(RetrievalResult(
                        content=chunk,
                        metadata={**doc.metadata, 'chunk_index': i, 'document_id': doc.id},
                        score=float(score),
                        retrieval_mode="chunks"
                    ))
        st.info(f"Found {len(results)} relevant chunks above threshold")
        return sorted(results, key=lambda x: x.score, reverse=True)[:top_k]
    
    def get_files_via_metadata(self, query: str) -> List[RetrievalResult]:
        results = []
        query_lower = query.lower()
        query_words = [word for word in query_lower.split() if len(word) > 2]
        st.info(f"Searching metadata for query words: {query_words}")
        for doc in self.documents:
            filename_lower = doc.metadata.get('filename', '').lower()
            filename_match = any(word in filename_lower for word in query_words)
            content_match = self._calculate_keyword_relevance(query, doc.content) > 0.1
            if filename_match or content_match:
                score = 0.9 if filename_match else self._calculate_keyword_relevance(query, doc.content)
                results.append(RetrievalResult(
                    content=doc.content[:2000] + "..." if len(doc.content) > 2000 else doc.content,
                    metadata=doc.metadata,
                    score=score,
                    retrieval_mode="files_via_metadata"
                ))
                st.info(f"Found match in {filename_lower} with score {score:.3f}")
        return results
    
    def get_files_via_content(self, query: str) -> List[RetrievalResult]:
        results = []
        st.info(f"Searching {len(self.documents)} documents by content")
        for doc in self.documents:
            if self.processor.embedder and doc.embedding is not None:
                try:
                    query_embedding = self.processor.embedder.encode(query)
                    score = cosine_similarity([query_embedding], [doc.embedding])[0][0]
                except:
                    score = self._calculate_keyword_relevance(query, doc.content)
            else:
                score = self._calculate_keyword_relevance(query, doc.content)
            if score > 0.1:
                results.append(RetrievalResult(
                    content=doc.content[:3000] + "..." if len(doc.content) > 3000 else doc.content,
                    metadata=doc.metadata,
                    score=float(score),
                    retrieval_mode="files_via_content"
                ))
                st.info(f"Document {doc.metadata.get('filename')} scored {score:.3f}")
        return sorted(results, key=lambda x: x.score, reverse=True)
    
    def _calculate_keyword_relevance(self, query: str, content: str) -> float:
        query_words = set(word.lower().strip('.,!?;:') for word in query.split() if len(word) > 2)
        content_words = set(word.lower().strip('.,!?;:') for word in content.split())
        if not query_words:
            return 0.0
        overlap = len(query_words.intersection(content_words))
        relevance = overlap / len(query_words)
        if query.lower() in content.lower():
            relevance += 0.3
        return min(relevance, 1.0)

class AgenticRetriever:
    def __init__(self, groq_api_key: str):
        self.groq_api_key = groq_api_key
        self.knowledge_bases: Dict[str, KnowledgeBase] = {}
    
    def call_groq_api(self, messages: list, max_retries: int = 3) -> str:
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {self.groq_api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "llama-3.3-70b-versatile",  # Updated to a known versatile model
            "messages": messages,
            "max_tokens": 1500,
            "temperature": 0.3
        }
        for attempt in range(max_retries):
            try:
                response = requests.post(url, headers=headers, json=payload, timeout=30)
                response.raise_for_status()
                return response.json()['choices'][0]['message']['content']
            except requests.exceptions.RequestException as e:
                if attempt == max_retries - 1:
                    raise Exception(f"Failed to call Groq API after {max_retries} attempts: {str(e)}")
                time.sleep(2 ** attempt)
    
    def add_knowledge_base(self, kb: KnowledgeBase):
        self.knowledge_bases[kb.name] = kb
        st.success(f"Added knowledge base: {kb.name}")
    
    def route_to_knowledge_base(self, query: str) -> List[str]:
        if not self.knowledge_bases:
            return []
        kb_descriptions = "\n".join([
            f"- {name}: {kb.description} ({len(kb.documents)} documents)" 
            for name, kb in self.knowledge_bases.items()
        ])
        messages = [
            {
                "role": "system",
                "content": "You are an intelligent routing agent. Given a user query and available knowledge bases, determine which knowledge base(s) would be most relevant to answer the query. Respond with only the knowledge base name(s), separated by commas if multiple."
            },
            {
                "role": "user",
                "content": f"Query: {query}\n\nAvailable Knowledge Bases:\n{kb_descriptions}\n\nWhich knowledge base(s) should be used to answer this query?"
            }
        ]
        try:
            response = self.call_groq_api(messages)
            selected_kbs = [kb.strip() for kb in response.split(',') if kb.strip() in self.knowledge_bases]
            return selected_kbs if selected_kbs else list(self.knowledge_bases.keys())
        except Exception as e:
            st.error(f"Error in KB routing: {str(e)}")
            return list(self.knowledge_bases.keys())
    
    def determine_retrieval_mode(self, query: str, kb_name: str) -> RetrievalMode:
        messages = [
            {
                "role": "system",
                "content": "You are a retrieval mode classifier. Given a query, determine the best retrieval strategy:\n\n- \"chunks\": For specific information queries that can be answered with document excerpts\n- \"files_via_metadata\": When the query mentions specific file names or asks about particular documents\n- \"files_via_content\": For broad topic queries needing comprehensive info from entire documents\n- \"web_search\": When the query requires external information not likely in the documents\n\nRespond with only one of these options: chunks, files_via_metadata, files_via_content, web_search"
            },
            {
                "role": "user",
                "content": f"Query: {query}\nKnowledge Base: {kb_name}\n\nBest retrieval mode:"
            }
        ]
        try:
            response = self.call_groq_api(messages).strip().lower()
            if "files_via_metadata" in response:
                return RetrievalMode.FILES_VIA_METADATA
            elif "files_via_content" in response:
                return RetrievalMode.FILES_VIA_CONTENT
            elif "web_search" in response:
                return RetrievalMode.WEB_SEARCH
            else:
                return RetrievalMode.CHUNKS
        except Exception as e:
            st.error(f"Error in mode determination: {str(e)}")
            return RetrievalMode.CHUNKS
    
    def perform_web_search(self, query: str) -> List[RetrievalResult]:
        # Mock web search (since real API integration is out of scope)
        st.info("Performing mock web search...")
        return [RetrievalResult(
            content=f"Mock web result for '{query}': This is a simulated response.",
            metadata={'source': 'web'},
            score=0.7,
            retrieval_mode="web_search"
        )]
    
    def evaluate_results(self, query: str, results: List[RetrievalResult]) -> bool:
        if not results:
            return False
        top_score = max(result.score for result in results)
        messages = [
            {
                "role": "system",
                "content": "You are a result evaluator. Given a query and retrieved content, determine if the results are sufficient to answer the query. Respond with 'yes' or 'no'."
            },
            {
                "role": "user",
                "content": f"Query: {query}\n\nRetrieved Content:\n{results[0].content[:1000]}\n\nAre these results sufficient?"
            }
        ]
        try:
            response = self.call_groq_api(messages).strip().lower()
            return response == 'yes'
        except Exception:
            return top_score > 0.5  # Fallback to score threshold
    
    def retrieve(self, query: str, top_k: int = 5) -> List[RetrievalResult]:
        if not self.knowledge_bases:
            st.warning("No knowledge bases available. Please upload some documents first.")
            return []
        
        selected_kbs = self.route_to_knowledge_base(query)
        st.info(f"🎯 Routed query to knowledge bases: {', '.join(selected_kbs)}")
        
        all_results = []
        attempted_modes = set()
        
        for kb_name in selected_kbs:
            if kb_name not in self.knowledge_bases:
                continue
            kb = self.knowledge_bases[kb_name]
            if not kb.documents:
                st.warning(f"No documents in knowledge base: {kb_name}")
                continue
            
            retrieval_mode = self.determine_retrieval_mode(query, kb_name)
            attempted_modes.add(retrieval_mode)
            st.info(f"📊 Using retrieval mode '{retrieval_mode.value}' for {kb_name}")
            
            if retrieval_mode == RetrievalMode.CHUNKS:
                results = kb.get_relevant_chunks(query, top_k)
            elif retrieval_mode == RetrievalMode.FILES_VIA_METADATA:
                results = kb.get_files_via_metadata(query)
            elif retrieval_mode == RetrievalMode.FILES_VIA_CONTENT:
                results = kb.get_files_via_content(query)
            elif retrieval_mode == RetrievalMode.WEB_SEARCH:
                results = self.perform_web_search(query)
            else:
                results = []
            
            for result in results:
                result.metadata['knowledge_base'] = kb_name if retrieval_mode != RetrievalMode.WEB_SEARCH else 'web'
            all_results.extend(results)
            st.info(f"Retrieved {len(results)} results from {kb_name}")
            
            # Iterative refinement
            if not self.evaluate_results(query, all_results):
                st.info("Results insufficient, attempting refinement...")
                remaining_modes = {RetrievalMode.CHUNKS, RetrievalMode.FILES_VIA_CONTENT, RetrievalMode.WEB_SEARCH} - attempted_modes
                if remaining_modes:
                    next_mode = next(iter(remaining_modes))
                    attempted_modes.add(next_mode)
                    st.info(f"Trying alternative mode: {next_mode.value}")
                    if next_mode == RetrievalMode.CHUNKS:
                        extra_results = kb.get_relevant_chunks(query, top_k)
                    elif next_mode == RetrievalMode.FILES_VIA_CONTENT:
                        extra_results = kb.get_files_via_content(query)
                    elif next_mode == RetrievalMode.WEB_SEARCH:
                        extra_results = self.perform_web_search(query)
                    else:
                        extra_results = []
                    for result in extra_results:
                        result.metadata['knowledge_base'] = kb_name if next_mode != RetrievalMode.WEB_SEARCH else 'web'
                    all_results.extend(extra_results)
        
        all_results.sort(key=lambda x: x.score, reverse=True)
        final_results = all_results[:top_k]
        st.info(f"Final results: {len(final_results)} items")
        return final_results
    
    def generate_response(self, query: str, retrieved_results: List[RetrievalResult]) -> str:
        if not retrieved_results:
            return "I couldn't find relevant information to answer your query. Please try rephrasing or upload more relevant documents."
        context = "\n\n".join([
            f"Source: {result.metadata.get('filename', result.metadata.get('source', 'Unknown'))} "
            f"(KB: {result.metadata.get('knowledge_base', 'Unknown')}, Mode: {result.retrieval_mode})\n{result.content}"
            for result in retrieved_results
        ])
        messages = [
            {
                "role": "system",
                "content": "You are a helpful assistant. Use the provided context to answer the user's question comprehensively and accurately. Always cite which documents you're referencing. If the context doesn't contain enough information, say so clearly."
            },
            {
                "role": "user",
                "content": f"Context:\n{context}\n\nQuestion: {query}\n\nAnswer:"
            }
        ]
        try:
            return self.call_groq_api(messages)
        except Exception as e:
            return f"Error generating response: {str(e)}. However, I found relevant content:\n\n{context[:1000]}..."

# Streamlit App
def main():
    st.set_page_config(
        page_title="Agentic Retrieval System",
        page_icon="🤖",
        layout="wide"
    )
    st.title("🤖 Agentic Retrieval System")
    st.markdown("*Upload your documents and test agentic retrieval!*")
    
    st.sidebar.header("🔧 Debug Info")
    with st.sidebar:
        st.header("⚙️ Configuration")
        groq_api_key = os.getenv("GROQ_API_KEY")
        if not groq_api_key:
            st.warning("Please set GROQ_API_KEY environment variable for full agentic features.")
        if 'retriever' in st.session_state:
            st.write(f"KBs: {len(st.session_state.retriever.knowledge_bases)}")
            for name, kb in st.session_state.retriever.knowledge_bases.items():
                st.write(f"- {name}: {len(kb.documents)} docs")
    
    if 'retriever' not in st.session_state:
        st.session_state.retriever = AgenticRetriever(groq_api_key or "dummy_key")
    if 'processor' not in st.session_state:
        st.session_state.processor = DocumentProcessor()
    
    tab1, tab2, tab3 = st.tabs(["📄 Document Upload", "💬 Query & Test", "📊 Knowledge Bases"])
    
    with tab1:
        st.header("📄 Upload Documents")
        col1, col2 = st.columns([2, 1])
        with col1:
            kb_option = st.radio("Knowledge Base Options:", ["Create new knowledge base", "Add to existing knowledge base"])
        if kb_option == "Create new knowledge base":
            kb_name = st.text_input("Knowledge Base Name", placeholder="e.g., Financial Reports")
            kb_description = st.text_area("Description", placeholder="e.g., Company financial reports and SEC filings")
        else:
            existing_kbs = list(st.session_state.retriever.knowledge_bases.keys())
            if existing_kbs:
                kb_name = st.selectbox("Select existing knowledge base:", existing_kbs)
                kb_description = st.session_state.retriever.knowledge_bases[kb_name].description
            else:
                st.warning("No existing knowledge bases. Please create a new one.")
                kb_name = ""
                kb_description = ""
        uploaded_files = st.file_uploader("Upload documents", type=['pdf', 'docx', 'txt'], accept_multiple_files=True)
        if st.button("🚀 Process Documents", type="primary"):
            if not kb_name or not kb_description:
                st.error("Please provide knowledge base name and description.")
            elif not uploaded_files:
                st.error("Please upload at least one document.")
            else:
                if kb_name not in st.session_state.retriever.knowledge_bases:
                    kb = KnowledgeBase(kb_name, kb_description)
                    st.session_state.retriever.add_knowledge_base(kb)
                else:
                    kb = st.session_state.retriever.knowledge_bases[kb_name]
                progress_bar = st.progress(0)
                successful_uploads = 0
                for i, uploaded_file in enumerate(uploaded_files):
                    st.info(f"Processing {uploaded_file.name}...")
                    uploaded_file.seek(0)
                    document = st.session_state.processor.process_uploaded_file(uploaded_file)
                    if document:
                        kb.add_document(document)
                        successful_uploads += 1
                    progress_bar.progress((i + 1) / len(uploaded_files))
                if successful_uploads > 0:
                    st.success(f"✅ Processed {successful_uploads}/{len(uploaded_files)} documents!")
                    st.balloons()
                else:
                    st.error("❌ No documents processed. Check your files.")
                st.rerun()
    
    with tab2:
        st.header("💬 Query & Test Interface")
        if not st.session_state.retriever.knowledge_bases:
            st.warning("⚠️ No knowledge bases available. Upload documents first.")
        else:
            st.info(f"Available KBs: {', '.join(st.session_state.retriever.knowledge_bases.keys())}")
            if not groq_api_key:
                st.warning("⚠️ Groq API key not set. LLM features disabled.")
            query = st.text_area("Enter your query:", placeholder="e.g., What are the main findings in the uploaded report?", height=100)
            col1, col2 = st.columns([1, 1])
            with col1:
                top_k = st.slider("Number of results", 1, 10, 5)
            with col2:
                use_llm = st.checkbox("Use LLM for response generation", value=bool(groq_api_key), disabled=not bool(groq_api_key))
            if st.button("🔍 Retrieve & Answer", type="primary"):
                if query:
                    with st.spinner("🤖 Performing agentic retrieval..."):
                        results = st.session_state.retriever.retrieve(query, top_k)
                        if results:
                            st.success("✅ Retrieval Complete!")
                            if use_llm and groq_api_key:
                                try:
                                    response = st.session_state.retriever.generate_response(query, results)
                                    st.subheader("🤖 Generated Response")
                                    st.write(response)
                                except Exception as e:
                                    st.error(f"Error generating LLM response: {e}")
                                    st.info("Showing retrieved content instead:")
                            st.subheader("📄 Retrieved Context")
                            for i, result in enumerate(results, 1):
                                with st.expander(f"Result {i} - {result.metadata.get('filename', 'Unknown')} ({result.retrieval_mode}) - Score: {result.score:.3f}"):
                                    st.write("**Content Preview:**")
                                    st.write(result.content[:500] + "..." if len(result.content) > 500 else result.content)
                                    st.write("**Metadata:**")
                                    st.json(result.metadata)
                        else:
                            st.warning("No relevant results found. Try a different query or upload more documents.")
                            st.subheader("🔍 Debug Information")
                            total_docs = sum(len(kb.documents) for kb in st.session_state.retriever.knowledge_bases.values())
                            total_chunks = sum(len(doc.chunks or []) for kb in st.session_state.retriever.knowledge_bases.values() for doc in kb.documents)
                            st.write(f"Total documents: {total_docs}")
                            st.write(f"Total chunks: {total_chunks}")
                            for kb_name, kb in st.session_state.retriever.knowledge_bases.items():
                                for doc in kb.documents[:1]:
                                    st.write(f"Sample from {doc.metadata.get('filename')}:")
                                    st.write(doc.content[:200] + "...")
                else:
                    st.warning("Please enter a query.")
    
    with tab3:
        st.header("📊 Knowledge Bases Overview")
        if not st.session_state.retriever.knowledge_bases:
            st.info("No knowledge bases created yet. Upload some documents to get started!")
        else:
            for kb_name, kb in st.session_state.retriever.knowledge_bases.items():
                with st.expander(f"🗃️ {kb_name} ({len(kb.documents)} documents)", expanded=True):
                    st.write(f"**Description:** {kb.description}")
                    if kb.documents:
                        st.write("**Documents:**")
                        for doc in kb.documents:
                            col1, col2, col3, col4 = st.columns([3, 1, 1, 1])
                            with col1:
                                st.write(f"📄 {doc.metadata.get('filename', doc.id)}")
                            with col2:
                                st.write(f"{doc.metadata.get('file_type', 'unknown').upper()}")
                            with col3:
                                st.write(f"{doc.metadata.get('size', 0):,} chars")
                            with col4:
                                st.write(f"{doc.metadata.get('chunk_count', 0)} chunks")
                        doc_ids = [doc.id for doc in kb.documents]
                        doc_dict = {doc.id: doc for doc in kb.documents}
                        if doc_ids:
                            selected_doc_id = st.selectbox("Select a document to preview:", options=doc_ids, format_func=lambda x: doc_dict[x].metadata.get('filename', x), key=f"preview_select_{kb_name}")
                            if selected_doc_id:
                                selected_doc = doc_dict[selected_doc_id]
                                st.subheader(f"Preview: {selected_doc.metadata.get('filename', selected_doc.id)}")
                                st.write(selected_doc.content[:500] + "..." if len(selected_doc.content) > 500 else selected_doc.content)
                    else:
                        st.write("*No documents uploaded yet.*")
                    if st.button(f"🗑️ Delete {kb_name}", key=f"delete_{kb_name}"):
                        del st.session_state.retriever.knowledge_bases[kb_name]
                        st.success(f"Deleted knowledge base: {kb_name}")
                        st.rerun()

if __name__ == "__main__":
    main()